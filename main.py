#!/usr/bin/env python3
"""
Enhanced Tkinter GUI Application for YouTube Transcript Extraction, Chunk Splitting,
Processing with Ollama, and Combining Processed Chunks into a Document.

Features:
  - Single-window interface with frame swapping (no popups)
  - Modern, colorful interface with styled ttk elements
  - Centered text, custom fonts, and background gradients
  - Animated title color cycle
  - Animated spinner during processing
  - Inline error and status messages (no messageboxes)
  - Responsive layout that adjusts on window resize, ensuring all controls remain visible
  - Footer bar for filename entry and buttons
  - Cool startup splash animation
  - "Typewriter" effect for chunk response text display
  - Automatic retry on transcript extraction failure
  - Temporary files are cleared after saving, canceling, or backing out
  - Cross-platform (Windows, macOS, Linux)
  - DOCX title feature with centered filename
  - Sequential processing with typewriter effect per chunk
  - Enhanced settings options
  - Improved settings UI with tabs
  - Typewriter speed control
"""

import os
import sys
import json
import time
import shutil
import threading
import requests
import itertools
from pathlib import Path
from tkinter import (
    Tk,
    Frame,
    Text,
    Scrollbar,
    StringVar,
    IntVar,
    BooleanVar,
    CENTER,
    BOTH,
    END,
    W,
    E,
    N,
    S,
    ttk,
)
from tkinter.ttk import (
    Progressbar,
    Combobox,
    Checkbutton as TCheckbutton,
    Style,
    Label as TLabel,
    Button as TButton,
    Entry as TEntry,
    Notebook,
)
from tkinter.filedialog import asksaveasfilename
from youtube_transcript_api import (
    YouTubeTranscriptApi,
    NoTranscriptFound,
    TranscriptsDisabled,
    VideoUnavailable,
)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -------------------------
# Helper: read_file_with_fallback
# -------------------------
def read_file_with_fallback(filepath):
    encodings = ["utf-8", "latin1", "iso-8859-1"]
    for enc in encodings:
        try:
            with open(filepath, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError(f"Could not decode file {filepath} with available encodings.")


# -------------------------
# Ollama API Helper Function
# -------------------------
def generate_response(prompt, model, host="http://localhost:11434", cancel_event=None):
    url = f"{host}/api/generate"
    payload = {"model": model, "prompt": prompt, "stream": False}
    headers = {"Content-Type": "application/json"}
    try:
        if cancel_event and cancel_event.is_set():
            return "[Generation cancelled]", None
        response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=30)
        if cancel_event and cancel_event.is_set():
            return "[Generation cancelled]", None
        response.raise_for_status()
        try:
            json_response = response.json()
        except json.JSONDecodeError:
            # Retry once if empty response
            time.sleep(1)
            response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=30)
            response.raise_for_status()
            json_response = response.json()
        generated_text = json_response.get("response", "").strip()
        return generated_text, json_response
    except Exception as e:
        msg = str(e)
        if "no element found" in msg:
            return "[Unable to fetch transcript or response is empty]", None
        return f"[Error processing chunk: {e}]", None


# -------------------------
# Configuration
# -------------------------
class Config:
    def __init__(self):
        self.base_dir = Path(__file__).parent
        self.config_file = self.base_dir / "config.json"
        self.output_dir = self.base_dir / "outputs"
        self.temp_dir = self.base_dir / "temp"
        self._init_directories()
        self.settings = self._load_config()

    def _init_directories(self):
        self.output_dir.mkdir(exist_ok=True)
        self.temp_dir.mkdir(exist_ok=True)
        (self.temp_dir / "yt_trans").mkdir(exist_ok=True)
        (self.temp_dir / "yt_chunks").mkdir(exist_ok=True)
        (self.temp_dir / "yt_pro").mkdir(exist_ok=True)

    def _load_config(self):
        defaults = {
            "chunk_size": 300,
            "chunk_overlap": 50,
            "ollama_model": "deepseek-r1",
            "processing_prompt": "Check and reformat the text for grammar, clarity, and proper structure.",
            "output_format": "docx",
            "skip_manual_name": False,
            "last_video_id": "",
            "inline_output_name": "",
            "include_docx_title": True,
            "title_font_size": 16,
            "custom_title": "",
            "retry_count": 3,
            "typewriter_speed": 2,  # ms per character
        }
        try:
            with open(self.config_file, "r") as f:
                return {**defaults, **json.load(f)}
        except FileNotFoundError:
            return defaults

    def save_config(self):
        with open(self.config_file, "w") as f:
            json.dump(self.settings, f, indent=2)

    def clean_temp(self):
        for subdir in ["yt_trans", "yt_chunks", "yt_pro"]:
            dir_path = self.temp_dir / subdir
            for item in dir_path.glob("*"):
                if item.is_file():
                    item.unlink()


# -------------------------
# Transcript Handling
# -------------------------
class TranscriptHandler:
    def __init__(self, config: Config):
        self.config = config
        self.config.clean_temp()

    def extract_and_save_transcript(self, video_url):
        retry_count = int(self.config.settings.get("retry_count", 3))
        retry_delay = 1  # seconds between retries
        
        for attempt in range(retry_count + 1):
            try:
                if "youtu.be" in video_url:
                    video_id = video_url.split("/")[-1]
                else:
                    video_id = video_url.split("v=")[-1].split("&")[0]
                transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
                transcript_text = "\n".join([entry["text"] for entry in transcript_list])
                trans_dir = self.config.temp_dir / "yt_trans"
                transcript_file = trans_dir / f"{video_id}_transcript.txt"
                transcript_file.write_text(transcript_text, encoding="utf-8")
                self.config.settings["last_video_id"] = video_id
                self.config.save_config()
                return transcript_file, video_id
            except (NoTranscriptFound, TranscriptsDisabled, VideoUnavailable):
                if attempt < retry_count:
                    time.sleep(retry_delay)
                    retry_delay *= 2  # Exponential backoff
                    continue
                raise RuntimeError("Transcript unavailable for this video.")
            except Exception as e:
                msg = str(e)
                if attempt < retry_count:
                    time.sleep(retry_delay)
                    retry_delay *= 2  # Exponential backoff
                    continue
                if "no element found" in msg:
                    raise RuntimeError("Unable to fetch transcript; it may be unavailable or malformed.")
                raise RuntimeError(f"Error extracting transcript: {e}")

    def split_transcript(self, transcript_file):
        try:
            chunk_size = int(self.config.settings.get("chunk_size", 300))
            chunk_overlap = int(self.config.settings.get("chunk_overlap", 50))
            content = transcript_file.read_text(encoding="utf-8")
            words = content.split()
            total_words = len(words)
            chunks_dir = self.config.temp_dir / "yt_chunks"
            chunk_files = []
            start = 0
            chunk_id = 1
            while start < total_words:
                end = min(start + chunk_size, total_words)
                chunk_text = " ".join(words[start:end])
                chunk_file = chunks_dir / f"chunk_{chunk_id}.txt"
                chunk_file.write_text(chunk_text, encoding="utf-8")
                chunk_files.append(chunk_file)
                start += chunk_size - chunk_overlap
                chunk_id += 1
            return chunk_files
        except Exception as e:
            raise RuntimeError(f"Error splitting transcript: {e}")

    def process_single_chunk(self, chunk_file, cancel_event=None):
        try:
            chunk_content = chunk_file.read_text(encoding="utf-8")
            processing_prompt = self.config.settings.get(
                "processing_prompt",
                "Check and reformat the text for grammar, clarity, and proper structure.",
            )
            combined_prompt = (
                f"Processing Instruction:\n{processing_prompt}\n\n"
                f"Apply the above instruction to the following text:\n{chunk_content}"
            )
            generated_text, _ = generate_response(
                combined_prompt,
                self.config.settings.get("ollama_model", "deepseek-r1"),
                cancel_event=cancel_event,
            )
            
            # Save the processed chunk to yt_pro folder
            output_dir = self.config.temp_dir / "yt_pro"
            output_dir.mkdir(exist_ok=True)  # Ensure directory exists
            output_file = output_dir / chunk_file.name
            output_file.write_text(generated_text, encoding="utf-8")
            
            return generated_text
        except Exception as e:
            return f"[Error processing chunk: {e}]"

    def combine_chunks_to_output(self, video_id, status_callback=None):
        processed_dir = self.config.temp_dir / "yt_pro"
        processed_files = sorted(processed_dir.glob("*.txt"))
        if not processed_files:
            if status_callback:
                status_callback("Error: No processed chunks to combine.", "error")
            return

        if self.config.settings.get("skip_manual_name", False):
            default_name = video_id
        else:
            default_name = self.config.settings.get("inline_output_name", "").strip() or video_id

        output_format = self.config.settings.get("output_format", "docx").lower()
        filetypes = [("DOCX file", "*.docx"), ("TXT file", "*.txt")]
        def_ext = f".{output_format}"
        save_path = asksaveasfilename(
            parent=None,
            initialfile=default_name,
            defaultextension=def_ext,
            filetypes=filetypes
        )
        if not save_path:
            if status_callback:
                status_callback("Save cancelled by user.", "error")
            return

        try:
            if save_path.lower().endswith(".txt"):
                combined_text = "\n\n".join([read_file_with_fallback(f) for f in processed_files])
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(combined_text)
            else:
                doc = Document()
                
                # Add title if enabled
                if self.config.settings.get("include_docx_title", True):
                    # Determine title text
                    custom_title = self.config.settings.get("custom_title", "").strip()
                    title_text = custom_title if custom_title else Path(save_path).stem
                    
                    # Add title paragraph
                    title_para = doc.add_paragraph()
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = title_para.add_run(title_text)
                    
                    # Set title font size
                    title_size = int(self.config.settings.get("title_font_size", 16))
                    title_run.font.size = Pt(title_size)
                    
                    # Add space after title
                    doc.add_paragraph()
                
                # Add processed content
                for file in processed_files:
                    content = read_file_with_fallback(file)
                    doc.add_paragraph(content)
                    doc.add_paragraph()
                doc.save(save_path)
            if status_callback:
                status_callback(f"Success: File saved at {save_path}", "success")
        except Exception as e:
            if status_callback:
                status_callback(f"Error saving file: {e}", "error")


# -------------------------
# Main Application GUI
# -------------------------
class YTTPApp:
    def __init__(self, root):
        self.root = root
        self.root.title("YTTP AI - Enhanced")
        self.root.geometry("800x600")
        self.root.minsize(700, 550)
        self.root.configure(bg="#1e1e2e")

        # Animation settings
        self.title_colors = ["#e0e0e0", "#b5b5e5", "#8a8acf", "#5f5fba"]
        self.title_color_index = 0

        # Initialize config and handler
        self.config = Config()
        self.handler = TranscriptHandler(self.config)

        # Style configuration
        self.style = Style()
        self.style.theme_use("clam")
        self.style.configure(
            "TLabel",
            background="#1e1e2e",
            foreground="#ffffff",
            font=("Segoe UI", 14),
        )
        self.style.configure(
            "Title.TLabel",
            font=("Segoe UI Semibold", 24),
            foreground=self.title_colors[0],
        )
        self.style.configure(
            "TButton",
            background="#3b3b5e",
            foreground="#ffffff",
            font=("Segoe UI", 12),
            padding=8,
        )
        self.style.map(
            "TButton",
            background=[("active", "#5a5a8a")],
            relief=[("pressed", "sunken"), ("!pressed", "raised")],
        )
        self.style.configure(
            "TEntry",
            fieldbackground="#2e2e3f",
            foreground="#ffffff",
            insertbackground="#ffffff",
            font=("Segoe UI", 12),
        )
        self.style.configure(
            "TCombobox",
            fieldbackground="#2e2e3f",
            foreground="#ffffff",
            font=("Segoe UI", 12),
        )
        self.style.configure(
            "TCheckbutton",
            background="#1e1e2e",
            foreground="#ffffff",
            font=("Segoe UI", 12),
        )
        self.style.configure(
            "Horizontal.TProgressbar",
            troughcolor="#2e2e3f",
            background="#6a6ab5",
            thickness=20,
        )
        # Tab style
        self.style.configure(
            "TNotebook.Tab",
            background="#2e2e3f",
            foreground="#ffffff",
            font=("Segoe UI", 10),
            padding=[10, 5],
        )
        self.style.map(
            "TNotebook.Tab",
            background=[("selected", "#4a4a6a")],
            expand=[("selected", [1, 1, 1, 0])],
        )

        # Main container frame
        self.container = Frame(self.root, bg="#1e1e2e")
        self.container.pack(fill=BOTH, expand=True)

        # Dictionary to hold different frames
        self.frames = {}

        # Initialize frames, including SplashFrame
        for F in (SplashFrame, MenuFrame, StartFrame, ProcessingFrame, SettingsFrame):
            frame = F(parent=self.container, controller=self)
            self.frames[F.__name__] = frame
            frame.grid(row=0, column=0, sticky=N + S + E + W)

        # Show the splash frame initially
        self.show_frame("SplashFrame")

        # Start title animation (for frames that show Title)
        self.animate_title()

        # Configure container grid to expand
        self.container.rowconfigure(0, weight=1)
        self.container.columnconfigure(0, weight=1)

    def show_frame(self, frame_name):
        frame = self.frames[frame_name]
        frame.tkraise()

    def animate_title(self):
        # Cycle through title colors for frames with "Title.TLabel"
        self.title_color_index = (self.title_color_index + 1) % len(self.title_colors)
        self.style.configure("Title.TLabel", foreground=self.title_colors[self.title_color_index])
        # Repeat every 500ms
        self.root.after(500, self.animate_title)

    def start_processing_thread(self, video_url):
        frame = self.frames["ProcessingFrame"]
        frame.start_processing(video_url)

    def combine_output(self):
        frame = self.frames["ProcessingFrame"]
        frame.combine_output()

    def exit_application(self):
        self.config.clean_temp()
        self.root.destroy()


# -------------------------
# Splash Frame
# -------------------------
class SplashFrame(Frame):
    def __init__(self, parent, controller: YTTPApp):
        super().__init__(parent, bg="#1e1e2e")
        self.controller = controller
        # Full-screen text area for splash
        self.splash_label = TLabel(self, text="", style="Title.TLabel", anchor=CENTER)
        self.splash_label.place(relx=0.5, rely=0.5, anchor=CENTER)
        self.text_to_type = "-- == YTTP AI == --"
        self.char_index = 0
        self.after(500, self.type_text)  # Start after short delay

    def type_text(self):
        if self.char_index < len(self.text_to_type):
            current_text = self.splash_label.cget("text") + self.text_to_type[self.char_index]
            self.splash_label.config(text=current_text)
            self.char_index += 1
            # Lightning speed typing
            self.after(50, self.type_text)
        else:
            # Hold for a moment then go to Menu
            self.after(1000, lambda: self.controller.show_frame("MenuFrame"))


# -------------------------
# Menu Frame
# -------------------------
class MenuFrame(Frame):
    def __init__(self, parent, controller: YTTPApp):
        super().__init__(parent, bg="#1e1e2e")
        self.controller = controller

        title_label = TLabel(self, text="-- == YTTP AI == --", style="Title.TLabel")
        title_label.pack(pady=(40, 30))

        buttons_frame = Frame(self, bg="#1e1e2e")
        buttons_frame.pack(expand=True)

        start_btn = TButton(
            buttons_frame,
            text="Start",
            width=20,
            command=lambda: controller.show_frame("StartFrame"),
        )
        start_btn.grid(row=0, column=0, padx=20, pady=20)

        output_btn = TButton(
            buttons_frame,
            text="Output",
            width=20,
            command=controller.combine_output,
        )
        output_btn.grid(row=0, column=1, padx=20, pady=20)

        settings_btn = TButton(
            buttons_frame,
            text="Settings",
            width=20,
            command=lambda: controller.show_frame("SettingsFrame"),
        )
        settings_btn.grid(row=1, column=0, padx=20, pady=20)

        exit_btn = TButton(
            buttons_frame,
            text="Exit",
            width=20,
            command=controller.exit_application,
        )
        exit_btn.grid(row=1, column=1, padx=20, pady=20)

        buttons_frame.columnconfigure(0, weight=1)
        buttons_frame.columnconfigure(1, weight=1)


# -------------------------
# Start Frame
# -------------------------
class StartFrame(Frame):
    def __init__(self, parent, controller: YTTPApp):
        super().__init__(parent, bg="#1e1e2e")
        self.controller = controller

        header = TLabel(self, text="Enter YouTube URL:", style="TLabel")
        header.pack(pady=(30, 10))

        self.url_var = StringVar()
        url_entry = TEntry(self, textvariable=self.url_var, width=50, justify=CENTER)
        url_entry.pack(pady=(0, 10))
        url_entry.focus_set()

        self.error_label = TLabel(self, text="", style="TLabel")
        self.error_label.configure(foreground="#ff7373")
        self.error_label.pack(pady=(0, 10))

        btn_frame = Frame(self, bg="#1e1e2e")
        btn_frame.pack(pady=10)

        submit_btn = TButton(
            btn_frame,
            text="Submit",
            width=15,
            command=self.on_submit,
        )
        submit_btn.grid(row=0, column=0, padx=10)

        back_btn = TButton(
            btn_frame,
            text="Back",
            width=15,
            command=self.back_to_menu,
        )
        back_btn.grid(row=0, column=1, padx=10)

        btn_frame.columnconfigure(0, weight=1)
        btn_frame.columnconfigure(1, weight=1)

    def on_submit(self):
        url = self.url_var.get().strip()
        if not url:
            self.error_label.config(text="Error: URL cannot be empty.")
            return
        if ("youtube.com" not in url) and ("youtu.be" not in url):
            self.error_label.config(text="Error: Invalid YouTube URL format.")
            return
        self.error_label.config(text="")
        self.controller.show_frame("ProcessingFrame")
        threading.Thread(target=self.controller.start_processing_thread, args=(url,), daemon=True).start()

    def back_to_menu(self):
        self.controller.config.clean_temp()
        self.controller.show_frame("MenuFrame")


# -------------------------
# Processing Frame
# -------------------------
class ProcessingFrame(Frame):
    def __init__(self, parent, controller: YTTPApp):
        super().__init__(parent, bg="#1e1e2e")
        self.controller = controller
        self.cancel_event = threading.Event()
        self.current_chunk_index = 0
        self.total_chunks = 0
        self.chunk_files = []
        self.video_id = ""
        self.display_text = ""
        self.display_index = 0

        # Status and spinner
        top_frame = Frame(self, bg="#1e1e2e")
        top_frame.pack(pady=(20, 5), fill="x")

        self.progress_label = TLabel(top_frame, text="Processing: 0/0", style="TLabel")
        self.progress_label.pack(side="left", padx=(20, 0))

        self.spinner_label = TLabel(top_frame, text="", style="TLabel", font=("Segoe UI", 16))
        self.spinner_label.pack(side="right", padx=(0, 20))

        self.progress_bar = Progressbar(self, orient="horizontal", mode="determinate", style="Horizontal.TProgressbar")
        self.progress_bar.pack(pady=(0, 15), fill="x", padx=20)

        # Text area with fixed height and scrollbar
        text_frame = Frame(self, bg="#2e2e3f")
        text_frame.pack(fill=BOTH, expand=True, padx=20, pady=(0, 5))
        scrollbar = Scrollbar(text_frame)
        scrollbar.pack(side="right", fill="y")
        self.response_text = Text(
            text_frame,
            wrap="word",
            yscrollcommand=scrollbar.set,
            bg="#2e2e3f",
            fg="#f0f0f0",
            insertbackground="#f0f0f0",
            bd=2,
            relief="groove",
            font=("Segoe UI", 11),
            height=15,  # Fixed height
        )
        self.response_text.pack(fill=BOTH, expand=True)
        scrollbar.config(command=self.response_text.yview)

        # Footer bar: filename entry and buttons
        footer_frame = Frame(self, bg="#1e1e2e", pady=5)
        footer_frame.pack(fill="x", side="bottom")

        TLabel(footer_frame, text="Filename:", style="TLabel").grid(row=0, column=0, padx=(20, 5), sticky=W)
        self.out_filename_var = StringVar(value=controller.config.settings.get("inline_output_name", ""))
        self.out_filename_entry = TEntry(footer_frame, textvariable=self.out_filename_var, width=25, justify=CENTER)
        self.out_filename_entry.grid(row=0, column=1, padx=(0, 10), sticky=W)

        combine_btn = TButton(
            footer_frame,
            text="Combine",
            width=12,
            command=self.combine_output,
        )
        combine_btn.grid(row=0, column=2, padx=10)

        cancel_btn = TButton(
            footer_frame,
            text="Cancel",
            width=12,
            command=self.cancel_processing,
        )
        cancel_btn.grid(row=0, column=3, padx=10)

        back_btn = TButton(
            footer_frame,
            text="Back",
            width=12,
            command=self.back_to_menu,
        )
        back_btn.grid(row=0, column=4, padx=10)

        self.status_label = TLabel(footer_frame, text="", style="TLabel")
        self.status_label.configure(font=("Segoe UI", 12))
        self.status_label.grid(row=1, column=0, columnspan=5, pady=(5, 0), sticky=W, padx=20)

        footer_frame.columnconfigure(1, weight=1)
        footer_frame.columnconfigure(2, weight=0)
        footer_frame.columnconfigure(3, weight=0)
        footer_frame.columnconfigure(4, weight=0)

    def start_processing(self, video_url):
        # Reset state
        self.cancel_event.clear()
        self.current_chunk_index = 0
        self.display_text = ""
        self.display_index = 0
        self.response_text.delete(1.0, END)
        self.progress_bar["value"] = 0
        self.progress_label.config(text="Processing: 0/0")
        self.status_label.config(text="", foreground="#ffffff")
        self.out_filename_var.set(self.controller.config.settings.get("last_video_id", ""))

        # Step 1: Extract transcript with retry loop
        while True:
            if self.cancel_event.is_set():
                self.status_label.config(text="Extraction cancelled.", foreground="#ff7373")
                self.controller.config.clean_temp()
                return
            try:
                transcript_file, video_id = self.controller.handler.extract_and_save_transcript(video_url)
                self.video_id = video_id
                break
            except RuntimeError as e:
                msg = str(e)
                if "Unable to fetch transcript" in msg:
                    self.status_label.config(text="Retrying transcript extraction...", foreground="#ff7373")
                    self.update_idletasks()
                    time.sleep(1)
                    continue
                else:
                    self.status_label.config(text=f"Error extracting transcript: {e}", foreground="#ff7373")
                    self.controller.config.clean_temp()
                    return

        # Step 2: Split transcript
        try:
            self.chunk_files = self.controller.handler.split_transcript(transcript_file)
            self.total_chunks = len(self.chunk_files)
            self.progress_bar["maximum"] = 100
            self.progress_label.config(text=f"Processing: 0/{self.total_chunks}")
        except Exception as e:
            self.status_label.config(text=f"Error splitting transcript: {e}", foreground="#ff7373")
            self.controller.config.clean_temp()
            return

        # Start processing first chunk
        self.process_next_chunk()

    def process_next_chunk(self):
        if self.current_chunk_index >= self.total_chunks or self.cancel_event.is_set():
            # Processing complete
            if not self.cancel_event.is_set():
                self.status_label.config(text="Processing complete. Enter filename and press Combine.", foreground="#b5e0a8")
                self.out_filename_var.set(self.video_id)
                self.controller.config.settings["inline_output_name"] = self.video_id
                self.controller.config.save_config()
            return

        # Update progress
        self.progress_label.config(text=f"Processing: {self.current_chunk_index+1}/{self.total_chunks}")
        self.progress_bar["value"] = int(((self.current_chunk_index+1) / self.total_chunks) * 100)
        
        # Animate spinner
        spinner = itertools.cycle(["◐", "◓", "◑", "◒"])
        self.animate_spinner(spinner)
        
        # Start processing this chunk in a separate thread
        chunk_file = self.chunk_files[self.current_chunk_index]
        threading.Thread(
            target=self.process_chunk_in_thread, 
            args=(chunk_file,),
            daemon=True
        ).start()

    def animate_spinner(self, spinner):
        if self.cancel_event.is_set():
            self.spinner_label.config(text="")
            return
            
        symbol = next(spinner)
        self.spinner_label.config(text=symbol)
        self.after(100, lambda: self.animate_spinner(spinner))

    def process_chunk_in_thread(self, chunk_file):
        try:
            # Process the chunk
            generated_text = self.controller.handler.process_single_chunk(
                chunk_file, 
                cancel_event=self.cancel_event
            )
            
            # Prepare to display
            header = f"\n--- Chunk {self.current_chunk_index+1} Response ---\n\n"
            self.display_text = header + generated_text
            self.display_index = 0
            
            # Start typewriter effect
            self.after(0, self.typewriter_effect)
        except Exception as e:
            self.after(0, lambda: self.status_label.config(
                text=f"Error processing chunk: {e}", 
                foreground="#ff7373"
            ))

    def typewriter_effect(self):
        if self.display_index < len(self.display_text) and not self.cancel_event.is_set():
            char = self.display_text[self.display_index]
            self.response_text.insert(END, char)
            self.response_text.see(END)
            self.display_index += 1
            
            # Get typewriter speed from settings
            speed = int(self.controller.config.settings.get("typewriter_speed", 2))
            self.after(speed, self.typewriter_effect)  # Use configurable delay
        else:
            # Move to next chunk after display completes
            self.spinner_label.config(text="")
            self.current_chunk_index += 1
            self.process_next_chunk()

    def combine_output(self):
        # Save inline output name to config
        self.controller.config.settings["inline_output_name"] = self.out_filename_var.get().strip()
        self.controller.config.save_config()

        def status_callback(msg, level):
            color = "#b5e0a8" if level == "success" else "#ff7373"
            self.status_label.config(text=msg, foreground=color)
            # After saving, clear temp
            self.controller.config.clean_temp()

        self.controller.handler.combine_chunks_to_output(self.video_id, status_callback=status_callback)

    def cancel_processing(self):
        self.cancel_event.set()
        self.status_label.config(text="Cancelling...", foreground="#ff7373")
        self.after(500, lambda: self._back_and_clear())

    def back_to_menu(self):
        self.cancel_event.set()
        self.controller.config.clean_temp()
        self.controller.show_frame("MenuFrame")

    def _back_and_clear(self):
        self.controller.config.clean_temp()
        self.controller.show_frame("MenuFrame")


# -------------------------
# Settings Frame with Tabs
# -------------------------
class SettingsFrame(Frame):
    def __init__(self, parent, controller: YTTPApp):
        super().__init__(parent, bg="#1e1e2e")
        self.controller = controller

        header = TLabel(self, text="Settings", style="Title.TLabel")
        header.pack(pady=(20, 10))

        # Create notebook (tabbed interface)
        notebook = Notebook(self, style="TNotebook")
        notebook.pack(fill=BOTH, expand=True, padx=20, pady=(0, 20))
        
        # Create tabs
        chunk_tab = Frame(notebook, bg="#1e1e2e")
        processing_tab = Frame(notebook, bg="#1e1e2e")
        output_tab = Frame(notebook, bg="#1e1e2e")
        
        notebook.add(chunk_tab, text="Chunk Settings")
        notebook.add(processing_tab, text="Processing Settings")
        notebook.add(output_tab, text="Output Settings")
        
        # Variables
        self.vars = {
            "chunk_size": IntVar(value=controller.config.settings.get("chunk_size", 300)),
            "chunk_overlap": IntVar(value=controller.config.settings.get("chunk_overlap", 50)),
            "retry_count": IntVar(value=controller.config.settings.get("retry_count", 3)),
            "ollama_model": StringVar(value=controller.config.settings.get("ollama_model", "deepseek-r1")),
            "processing_prompt": StringVar(value=controller.config.settings.get("processing_prompt", "")),
            "output_format": StringVar(value=controller.config.settings.get("output_format", "docx")),
            "skip_manual_name": BooleanVar(value=controller.config.settings.get("skip_manual_name", False)),
            "include_docx_title": BooleanVar(value=controller.config.settings.get("include_docx_title", True)),
            "title_font_size": IntVar(value=controller.config.settings.get("title_font_size", 16)),
            "custom_title": StringVar(value=controller.config.settings.get("custom_title", "")),
            "typewriter_speed": IntVar(value=controller.config.settings.get("typewriter_speed", 2)),
        }
        
        # Configure grid for tabs
        for tab in [chunk_tab, processing_tab, output_tab]:
            tab.columnconfigure(0, weight=1)
            tab.columnconfigure(1, weight=1)
            tab.columnconfigure(2, weight=1)
            tab.rowconfigure(0, weight=1)
        
        # Chunk Settings Tab
        self.create_chunk_settings(chunk_tab)
        
        # Processing Settings Tab
        self.create_processing_settings(processing_tab)
        
        # Output Settings Tab
        self.create_output_settings(output_tab)
        
        # Status label and buttons
        self.status_label = TLabel(self, text="", style="TLabel")
        self.status_label.configure(font=("Segoe UI", 12))
        self.status_label.pack(pady=(0, 10))

        btn_frame = Frame(self, bg="#1e1e2e")
        btn_frame.pack(pady=20)

        save_btn = TButton(
            btn_frame,
            text="Save Settings",
            width=20,
            command=self.on_save,
        )
        save_btn.grid(row=0, column=0, padx=10)

        back_btn = TButton(
            btn_frame,
            text="Back to Menu",
            width=20,
            command=self.back_to_menu,
        )
        back_btn.grid(row=0, column=1, padx=10)

    def create_chunk_settings(self, parent):
        # Chunk size
        TLabel(parent, text="Chunk Size:", style="TLabel").grid(row=0, column=0, padx=10, pady=10, sticky=W)
        chunk_size_entry = TEntry(parent, textvariable=self.vars["chunk_size"], width=10)
        chunk_size_entry.grid(row=0, column=1, padx=10, pady=10, sticky=W)
        
        # Chunk overlap
        TLabel(parent, text="Chunk Overlap:", style="TLabel").grid(row=1, column=0, padx=10, pady=10, sticky=W)
        chunk_overlap_entry = TEntry(parent, textvariable=self.vars["chunk_overlap"], width=10)
        chunk_overlap_entry.grid(row=1, column=1, padx=10, pady=10, sticky=W)
        
        # Retry count
        TLabel(parent, text="Retry Count:", style="TLabel").grid(row=2, column=0, padx=10, pady=10, sticky=W)
        retry_entry = TEntry(parent, textvariable=self.vars["retry_count"], width=10)
        retry_entry.grid(row=2, column=1, padx=10, pady=10, sticky=W)
        
        # Description
        desc = TLabel(parent, text="Chunk size and overlap are in words. Retry count is for transcript extraction.", 
                     style="TLabel", foreground="#a0a0c0", font=("Segoe UI", 10))
        desc.grid(row=3, column=0, columnspan=3, padx=10, pady=(20, 10), sticky=W)

    def create_processing_settings(self, parent):
        # Ollama model
        TLabel(parent, text="Ollama Model:", style="TLabel").grid(row=0, column=0, padx=10, pady=10, sticky=W)
        ollama_model_entry = TEntry(parent, textvariable=self.vars["ollama_model"], width=30)
        ollama_model_entry.grid(row=0, column=1, columnspan=2, padx=10, pady=10, sticky=W)
        
        # Processing prompt
        TLabel(parent, text="Processing Prompt:", style="TLabel").grid(row=1, column=0, padx=10, pady=10, sticky=W)
        processing_prompt_entry = Text(
            parent, 
            width=40, 
            height=6,
            bg="#2e2e3f",
            fg="#f0f0f0",
            insertbackground="#f0f0f0",
            font=("Segoe UI", 10),
            wrap="word"
        )
        processing_prompt_entry.grid(row=1, column=1, columnspan=2, padx=10, pady=10, sticky=W)
        processing_prompt_entry.insert("1.0", self.vars["processing_prompt"].get())
        self.processing_prompt_widget = processing_prompt_entry
        
        # Description
        desc = TLabel(parent, text="This prompt will be sent to Ollama with each chunk of text.", 
                     style="TLabel", foreground="#a0a0c0", font=("Segoe UI", 10))
        desc.grid(row=2, column=0, columnspan=3, padx=10, pady=(20, 10), sticky=W)

    def create_output_settings(self, parent):
        # Output format
        TLabel(parent, text="Output Format:", style="TLabel").grid(row=0, column=0, padx=10, pady=10, sticky=W)
        output_format_combo = Combobox(
            parent,
            textvariable=self.vars["output_format"],
            values=["docx", "txt"],
            state="readonly",
            width=10,
        )
        output_format_combo.grid(row=0, column=1, padx=10, pady=10, sticky=W)
        
        # Skip manual naming
        skip_check = TCheckbutton(
            parent,
            text="Skip Manual Naming",
            variable=self.vars["skip_manual_name"],
            style="TCheckbutton",
        )
        skip_check.grid(row=1, column=0, padx=10, pady=10, sticky=W)
        
        # Include DOCX title
        title_check = TCheckbutton(
            parent,
            text="Include Title in DOCX",
            variable=self.vars["include_docx_title"],
            style="TCheckbutton",
        )
        title_check.grid(row=1, column=1, padx=10, pady=10, sticky=W)
        
        # Title font size
        TLabel(parent, text="Title Font Size:", style="TLabel").grid(row=2, column=0, padx=10, pady=10, sticky=W)
        title_size_entry = TEntry(parent, textvariable=self.vars["title_font_size"], width=10)
        title_size_entry.grid(row=2, column=1, padx=10, pady=10, sticky=W)
        
        # Custom title
        TLabel(parent, text="Custom Title:", style="TLabel").grid(row=3, column=0, padx=10, pady=10, sticky=W)
        custom_title_entry = TEntry(parent, textvariable=self.vars["custom_title"], width=30)
        custom_title_entry.grid(row=3, column=1, columnspan=2, padx=10, pady=10, sticky=W)
        
        # Typewriter speed
        TLabel(parent, text="Typewriter Speed (ms):", style="TLabel").grid(row=4, column=0, padx=10, pady=10, sticky=W)
        speed_entry = TEntry(parent, textvariable=self.vars["typewriter_speed"], width=10)
        speed_entry.grid(row=4, column=1, padx=10, pady=10, sticky=W)
        
        # Description
        desc = TLabel(parent, text="Leave custom title blank to use filename as title.", 
                     style="TLabel", foreground="#a0a0c0", font=("Segoe UI", 10))
        desc.grid(row=5, column=0, columnspan=3, padx=10, pady=(20, 10), sticky=W)

    def on_save(self):
        try:
            self.controller.config.settings["chunk_size"] = int(self.vars["chunk_size"].get())
            self.controller.config.settings["chunk_overlap"] = int(self.vars["chunk_overlap"].get())
            self.controller.config.settings["retry_count"] = int(self.vars["retry_count"].get())
            self.controller.config.settings["ollama_model"] = self.vars["ollama_model"].get().strip()
            
            # Get processing prompt from text widget
            proc_prompt = self.processing_prompt_widget.get("1.0", END).strip()
            if proc_prompt:
                self.controller.config.settings["processing_prompt"] = proc_prompt
                
            self.controller.config.settings["output_format"] = self.vars["output_format"].get()
            self.controller.config.settings["skip_manual_name"] = self.vars["skip_manual_name"].get()
            self.controller.config.settings["include_docx_title"] = self.vars["include_docx_title"].get()
            self.controller.config.settings["title_font_size"] = int(self.vars["title_font_size"].get())
            self.controller.config.settings["custom_title"] = self.vars["custom_title"].get().strip()
            self.controller.config.settings["typewriter_speed"] = int(self.vars["typewriter_speed"].get())
            
            self.controller.config.save_config()
            self.status_label.config(text="Settings saved successfully.", foreground="#b5e0a8")
            self.controller.config.clean_temp()
        except ValueError:
            self.status_label.config(text="Error: Numeric values must be integers.", foreground="#ff7373")

    def back_to_menu(self):
        self.controller.config.clean_temp()
        self.controller.show_frame("MenuFrame")


if __name__ == "__main__":
    root = Tk()
    app = YTTPApp(root)
    root.mainloop()